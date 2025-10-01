import os
import PyPDF2
import docx
import re


class Readfile():
    """
    Класс получает файл, проверяет его тип. Если тип подходит, файл открывается и текст из него сохраняется в БД (будет осхраняться)
    """

    def __init__(self, name):
        self.name = name
        self.allowed_extensions = {'.txt', '.pdf', '.doc', '.docx'}
        self.content = None

    def _check_file_type(self):
        _, extension = os.path.splitext(self.name)
        if extension.lower() not in self.allowed_extensions:
            raise ValueError(
                f"Неподдерживаемый тип файла: {extension}. Разрешены: {', '.join(self.allowed_extensions)}")

        if not os.path.exists(self.name):
            raise FileNotFoundError(f"Файл {self.name} не найден")

        return extension.lower()

    def clear(self):

        patterns = [
            r'Рис\.\s*\d+\.\s*[^\n]+',  # Рис. 1. Описание
            r'Рисунок\s*\d+\.\s*[^\n]+',  # Рисунок 1. Описание
            r'Fig\.\s*\d+\.\s*[^\n]+',  # Fig. 1. Description
            r'Figure\s*\d+\.\s*[^\n]+',  # Figure 1. Description
            r'РИС\.\s*\d+\.\s*[^\n]+',  # РИС. 1. ОПИСАНИЕ
        ]

        cleaned_text = self.content
        for pattern in patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE)

        # Очистка результата
        cleaned_text = re.sub(r'\n\s*\n', '\n', cleaned_text)
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        cleaned_text = cleaned_text.strip()

        self.content = cleaned_text

    def read(self):
        try:
            extension = self._check_file_type()

            if extension == '.txt':
                self._read_txt()
            elif extension == '.pdf':
                self._read_pdf()
            elif extension in ['.doc', '.docx']:
                self._read_docx()



        except Exception as e:
            raise Exception(f"Ошибка при чтении файла: {str(e)}")

    def _read_txt(self):
        with open(self.name, 'r', encoding='utf-8') as file:
            self.content = file.read()

    def _read_pdf(self):
        try:
            with open(self.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                self.content = text
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")

    def _read_docx(self):
        """Читает DOC/DOCX файл"""
        try:
            doc = docx.Document(self.name)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            for table in doc.tables:
                text += self._read_docx_table(table)
            self.content = text
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX: {str(e)}")

    def _read_docx_table(self, table):
        """Извлекает текст из таблицы DOCX"""
        table_text = ""

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Рекурсивно обрабатываем вложенные таблицы
                if cell.tables:
                    for nested_table in cell.tables:
                        row_text.append(self._read_docx_table(nested_table))
                else:
                    # Обрабатываем текст ячейки
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text += paragraph.text + " "
                    row_text.append(cell_text.strip())

            # Фильтруем пустые ячейки и объединяем строку
            filtered_row = [cell for cell in row_text if cell]
            if filtered_row:
                table_text += " | ".join(filtered_row) + '\n'

        return table_text

    def save_data(self, output_file=None):
        """Сохраняет прочитанные данные в файл"""
        if self.content is None:
            raise ValueError("Сначала необходимо прочитать файл с помощью метода read()")

        if output_file is None:
            output_file = f"output_{os.path.basename(self.name)}.txt"

        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(self.content)

        return output_file

    def get_content(self):
        """Возвращает прочитанный контент"""
        return self.content

    def get_len(self):
        return len(self.content)


# Пример использования
if __name__ == "__main__":
    try:
        # Пример с текстовым файлом
        txt_reader = Readfile("tokarnaya.docx")
        txt_reader.read()
        txt_reader.clear()
        print("Текст из TXT файла:")
        print(txt_reader.get_content()[200:300] + "...")
        print()

        # Сохранение данных
        txt_reader.save_data("saved_txt.txt")

    except Exception as e:
        print(f"Ошибка: {e}")